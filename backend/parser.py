import pdfplumber
from docx import Document
from typing import Dict, Optional
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

async def parse_resume(file_path: str) -> Dict[str, any]:
    """Parse resume file (PDF, DOCX, or TXT) and extract text"""
    try:
        file_ext = Path(file_path).suffix.lower()
        text = ""
        
        if file_ext == ".pdf":
            text = await parse_pdf(file_path)
        elif file_ext == ".docx":
            text = await parse_docx(file_path)
        elif file_ext == ".txt":
            text = await parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Clean and normalize text
        text = clean_text(text)
        
        logger.info(f"Parsed {file_path}: {len(text)} characters")
        return {
            "text": text,
            "file_type": file_ext[1:],  # Remove dot
            "char_count": len(text)
        }
    except Exception as e:
        logger.error(f"Error parsing file {file_path}: {e}")
        raise

async def parse_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}")
        raise
    return text

async def parse_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        raise
    return text

async def parse_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
    except Exception as e:
        logger.error(f"Error parsing TXT: {e}")
        raise
    return text

def clean_text(text: str) -> str:
    """Clean and normalize extracted text"""
    # Remove excessive whitespace
    import re
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)]', '', text)
    # Strip leading/trailing whitespace
    text = text.strip()
    return text

